from docx import Document
import os
import json


# 读取 docx 文件，并返回，json格式，方便后续处理，json要返回要有自号，字体，格式，大小等、
def docx_to_json(docx_path):
    doc = Document(docx_path)
    data = []

    for para in doc.paragraphs:
        para_data = {
            'text': para.text if para.text else None,
            'style': para.style.name,
            'font': para.runs[0].font.name if para.runs else None,
            'size': para.runs[0].font.size.pt if para.runs and para.runs[0].font.size else None,
            'bold': para.runs[0].font.bold if para.runs else None,
            'italic': para.runs[0].font.italic if para.runs else None,
            'underline': para.runs[0].font.underline if para.runs else None
        }
        if para_data["text"] is None or para_data["text"].replace('\n', '').strip().replace(" ", "").replace("\t",
                                                                                                             "") == "":
            continue
        data.append(para_data)
        # print(para_data)
    # 将数据转换为JSON格式
    return json.dumps(data, ensure_ascii=False, indent=4)


# 读取docx文件，并返回纯文本，去除格式、换行符等
def docx_to_text(filename):
    doc = Document(filename)
    text = [paragraph.text for paragraph in doc.paragraphs]
    # 合并所有段落并去除多余的空格和换行符
    return ' '.join(text).replace('\n', '').strip().replace(" ", "").replace("\t", "")
    # return ' '.join(text).strip()


# 遍历指定目录内docx文件，并输出json文件
def all_docx_file_to_json(path):
    for root, dirs, files in os.walk(path):
        for file in files:
            if not file.endswith('.docx'):
                continue
            file_name = os.path.join(root, file)
            print('正在处理：', file_name)
            with open(file_name, 'rb') as f:
                try:
                    data = docx_to_json(f)
                    new_file_name = os.path.splitext(file)[0] + '.json'
                    new_file_name = os.path.join(root, new_file_name)
                    with open(new_file_name, 'w', encoding='utf-8') as w:
                        w.write(data)
                except Exception as e:
                    with open("error.log", 'a', encoding='utf-8') as w:
                        w.write('处理失败：' + file_name + ' 错误信息：' + str(e))
                    print('处理失败：', file_name, '错误信息：', str(e))
                print('处理完成')


# 测试
# if __name__ == '__main__':
#     path = r"法律条文"
#     docxfile_to_json(path)

# 遍历指定目录内docx文件，并输出txt文件，把输出的文件放入docx文件所在目录里的txt目录内，
def all_docx_file_to_txt(path):
    for root, dirs, files in os.walk(path):
        for file in files:
            if not file.endswith('.docx'):
                continue
            file_name = os.path.join(root, file)
            print('正在处理：', file_name)
            with open(file_name, 'rb') as f:
                try:
                    text = docx_to_text(f)
                    # 检测目录内是否有txt目录
                    if not os.path.exists(os.path.join(root, 'txt')):
                        os.mkdir(os.path.join(root, 'txt'))
                        print('创建txt目录')
                    new_file_name = os.path.splitext(file)[0] + '.txt'
                    new_file_name = os.path.join(root, "txt\\" + new_file_name)
                    with open(new_file_name, 'w', encoding='utf-8') as w:
                        w.write(text)
                except Exception as e:
                    with open("error.log", 'a', encoding='utf-8') as w:
                        w.write('处理失败：' + file_name + '错误信息：' + str(e))
                    print('处理失败：', file_name, '错误信息：', str(e))
    print('处理完成')


#
# 测试
if __name__ == '__main__':
    path = r"files/demo_2.docx"
    aaaa = docx_to_text(path)
    print(aaaa)
